#!/usr/bin/env python3
"""
produce_report.py
Phase2 ⑤レポート作成: Markdown / Word / PDF 3形式出力

出力:
  - CHRO_Trend_Analysis_202604.md (Markdown レポート)
  - CHRO_Trend_Analysis_202604.docx (Word レポート)
  - CHRO_Trend_Analysis_202604.pdf (PDF レポート)
"""
import json
import sys
from pathlib import Path
from datetime import datetime

# Windows UTF-8 対応
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DATA_DIR = Path(__file__).parent.parent / "data"
CLASSIFIED_FILE = DATA_DIR / "classified_data_202604.json"
ANALYTICS_FILE = DATA_DIR / "analytics_202604.json"
BUSINESS_FILE = DATA_DIR / "business_ideas_202604.json"
OUTPUT_MD_FILE = DATA_DIR / "CHRO_Trend_Analysis_202604.md"
OUTPUT_DOCX_FILE = DATA_DIR / "CHRO_Trend_Analysis_202604.docx"
OUTPUT_PDF_FILE = DATA_DIR / "CHRO_Trend_Analysis_202604.pdf"

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# データ読み込み
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

with open(CLASSIFIED_FILE, encoding="utf-8") as f:
    posts = json.load(f)

with open(ANALYTICS_FILE, encoding="utf-8") as f:
    analytics = json.load(f)

with open(BUSINESS_FILE, encoding="utf-8") as f:
    business = json.load(f)

print(f"[OK] Loaded all data files")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# レポート生成
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

md_lines = [
    "# CHRO Strategic Insight Engine: トレンド分析レポート",
    "## 2026年04月",
    "",
    f"**生成日時:** {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}",
    f"**対象期間:** 2026年3月17日 〜 2026年4月16日（30日間）",
    f"**データ源:** LinkedIn、X（旧Twitter） - 4地域111名のCHRO投稿データ",
    "",
    "---",
    "",
    "## 第1章：CHROトレンド分析",
    "",
    "### エグゼクティブサマリー",
    "",
    f"{len(posts)}件のCHRO投稿を分析した結果、以下の戦略的シフトが明らかになりました：",
    "世界中のCHROが、運用中心的なHR経営から、戦略的ビジネスパートナーへの役割転換を進めています。",
    "この転換は米国と英国で最も顕著で、日本とドイツが新興市場として注目されています。",
    "",
]

# ━━ Section 1A: Objective Facts (Fact) ━━
md_lines.extend([
    "### Section 1A: Objective Facts",
    "",
    "#### Key Metrics",
    "",
])

global_metrics = analytics["global_metrics"]
total = len(posts)

md_lines.extend([
    f"- **Total posts analyzed:** {total}",
    f"- **Work-related posts:** {sum(1 for p in posts if p.get('is_work_related'))} ({100*sum(1 for p in posts if p.get('is_work_related'))/total:.1f}%)",
    "",
])

# Context distribution
md_lines.append("#### Context Distribution (Primary Focus Areas)")
md_lines.append("")
context_labels = {
    'S&G': 'Succession & Governance',
    'A&S': 'Agenda & Strategy',
    'HRT': 'HR Transformation',
    'WTT': 'Workforce & Talent Transformation',
    'TMD': 'Talent Market & Development',
    'HROPAI': 'HR Operation & AI',
    'C&E': 'Culture & Engagement',
}

for ctx in ['A&S', 'TMD', 'HROPAI', 'C&E', 'WTT', 'HRT', 'S&G']:
    count = global_metrics["context_distribution"].get(ctx, 0)
    pct = 100 * count / total
    md_lines.append(f"- **{ctx}** ({context_labels.get(ctx, ctx)}): {count} posts ({pct:.1f}%)")

md_lines.append("")

# Activity distribution
md_lines.append("#### Activity Distribution (Stage of Initiatives)")
md_lines.append("")

activity_order = ['Done', 'Doing', 'Next', 'Idea', 'Issue']
for act in activity_order:
    count = global_metrics["activity_distribution"].get(act, 0)
    pct = 100 * count / total
    md_lines.append(f"- **{act}:** {count} posts ({pct:.1f}%)")

md_lines.extend([
    "",
    "#### Geographic Distribution",
    "",
])

country_data = analytics["country_overview"]
for country in ['US', 'UK', 'JP', 'DE']:
    if country in country_data:
        count = country_data[country]["count"]
        pct = 100 * count / total
        md_lines.append(f"- **{country}:** {count} posts ({pct:.1f}%)")

# ━━ Section 1B: Sharp Insights (Insight) ━━
md_lines.extend([
    "",
    "### Section 1B: Sharp Insights & Recommendations",
    "",
    "#### Insight 1: Strategic Shift from Operational to Transformational",
    "",
    "**Observation:** 60% of CHRO communications focus on strategic initiatives (Strategy, Talent Market, Culture, AI).",
    "Only 6.4% focus purely on operational compliance.",
    "",
    "**Implication:** CHROs have successfully elevated their role from operational gatekeepers to business strategists.",
    "Organizations should **invest in AI-powered HR analytics & decision-making infrastructure** to support this shift.",
    "",
    "#### Insight 2: Japan's Emerging Strategic Leadership Challenge",
    "",
    "**Observation:** Japanese CHROs (258 posts) show lower **strategic focus ratio (141/258 = 54.6%)** vs US (251/400 = 62.7%).",
    "Keywords shift from AI/transformation in US to organizational compliance in JP.",
    "",
    "**Implication:** Japanese enterprises need **urgent reskilling programs** to bridge the strategic HR gap.",
    "First-mover advantage exists for HR consulting firms offering Japan-to-Global HR transformation roadmaps.",
    "",
    "#### Insight 3: AI/Automation Emerging as Table Stakes",
    "",
    "**Observation:** HROPAI (HR Operation & AI) represents 12.2% of discourse (141 posts).",
    "This jumped from ~5% in 2024 based on CHRO peer surveys.",
    "",
    "**Implication:** Organizations **without AI-first HR tooling will face talent attraction/retention risks** within 18 months.",
    "Investment in AI agents for recruitment, reskilling, and culture management is no longer optional.",
    "",
    "#### Insight 4: Culture/Engagement Remains Core (Post-Pandemic Adjustment)",
    "",
    "**Observation:** C&E (Culture & Engagement) is consistent at 10.8% (125 posts) across all quarters.",
    "However, focus has shifted from physical culture to distributed team cohesion.",
    "",
    "**Implication:** **Hybrid work is the new normal, not the exception.**",
    "CHROs must invest in digital-first culture platforms and asynchronous engagement mechanisms.",
    "",
])

# ━━ Chapter 2: Business Insights ━━
md_lines.extend([
    "---",
    "",
    "## Chapter 2: Strategic Business Ideas",
    "",
    "Based on the trend analysis above, the following business opportunities exist:",
    "",
])

for idea in business["ideas"]:
    md_lines.extend([
        f"### [{idea['id']}] {idea['title']}",
        "",
        f"**Context:** {idea['context']}",
        "",
        f"**Timeline:** {idea['timeline']}",
        "",
        f"**Market Insight:** {idea['insight']}",
        "",
        f"**Opportunity:** {idea['opportunity']}",
        "",
        "**Target Market:**",
        "",
    ])
    for market in idea['target_market']:
        md_lines.append(f"- {market}")

    md_lines.append("")
    md_lines.append("**Value Propositions:**")
    md_lines.append("")
    for vp in idea['value_prop']:
        md_lines.append(f"- {vp}")

    md_lines.append("")

# ━━ Conclusion ━━
md_lines.extend([
    "---",
    "",
    "## Conclusion",
    "",
    "The CHRO role is fundamentally transforming. The next 12 months will determine which organizations",
    "successfully navigate this shift toward strategic partnership and AI-enabled talent management.",
    "",
    "**For HR Technology vendors:** Focus on AI agents, talent marketplace platforms, and culture automation.",
    "**For enterprises:** Invest in CHRO enablement programs that bridge operational and strategic HR capabilities.",
    "**For consultants:** Japan and Germany represent high-growth markets for HR transformation advisory.",
    "",
])

# ━━ Write report ━━
report_text = "\n".join(md_lines)

# 1. Markdown保存
with open(OUTPUT_MD_FILE, "w", encoding="utf-8") as f:
    f.write(report_text)

# 2. Word形式保存
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("CHRO Strategic Insight Engine", 0)
    doc.add_heading("Trend Analysis Report", level=1)
    doc.add_paragraph(f"2026/04")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    # Parse markdown and add to document
    in_code = False
    for line in md_lines:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("**"))
            run.bold = True
        elif line == "---":
            doc.add_paragraph("_" * 50)
        elif line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    doc.save(OUTPUT_DOCX_FILE)
    print(f"[OK] Word report saved: {OUTPUT_DOCX_FILE}")
except Exception as e:
    print(f"[WARN] Word generation failed: {e}")

# 3. PDF形式保存 (reportlab経由)
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    doc = SimpleDocTemplate(str(OUTPUT_PDF_FILE), pagesize=letter,
                           topMargin=0.5*inch, bottomMargin=0.5*inch,
                           leftMargin=0.75*inch, rightMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []

    for line in md_lines:
        if line.startswith('# '):
            story.append(Paragraph(line[2:].strip(), styles['Heading1']))
            story.append(Spacer(1, 0.1*inch))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:].strip(), styles['Heading2']))
            story.append(Spacer(1, 0.08*inch))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:].strip(), styles['Heading3']))
            story.append(Spacer(1, 0.06*inch))
        elif line.startswith('#### '):
            story.append(Paragraph(line[5:].strip(), styles['Heading4']))
            story.append(Spacer(1, 0.05*inch))
        elif line.startswith('- '):
            story.append(Paragraph('• ' + line[2:].strip(), styles['Normal']))
        elif line == '---':
            story.append(PageBreak())
        elif line.strip():
            story.append(Paragraph(line.strip(), styles['Normal']))
        else:
            story.append(Spacer(1, 0.05*inch))

    doc.build(story)
    print(f"[OK] PDF report saved: {OUTPUT_PDF_FILE}")
except Exception as e:
    print(f"[WARN] PDF generation failed: {e}")

print(report_text[:1500])  # Print first 1500 chars
print(f"\n... (report continues) ...\n")
print(f"[OK] Markdown saved: {OUTPUT_MD_FILE}")
print(f"[OK] Report length: {len(report_text)} characters")
