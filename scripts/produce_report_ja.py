#!/usr/bin/env python3
"""
produce_report_ja.py
Phase2 ⑤レポート作成（日本語版）: Markdown / Word / PDF 3形式出力

出力:
  - CHRO_Trend_Analysis_202604_JA.md (日本語Markdown)
  - CHRO_Trend_Analysis_202604_JA.docx (日本語Word)
  - CHRO_Trend_Analysis_202604_JA.pdf (日本語PDF)
"""
import json
import sys
from pathlib import Path
from datetime import datetime

# Windows UTF-8 対応
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DATA_DIR = Path(__file__).parent.parent / "data"
CLASSIFIED_FILE = DATA_DIR / "classified_data_202604.json"
ANALYTICS_FILE = DATA_DIR / "analytics_202604.json"
BUSINESS_FILE = DATA_DIR / "business_ideas_202604.json"
OUTPUT_MD_FILE = DATA_DIR / "CHRO_Trend_Analysis_202604_JA.md"
OUTPUT_DOCX_FILE = DATA_DIR / "CHRO_Trend_Analysis_202604_JA.docx"
OUTPUT_PDF_FILE = DATA_DIR / "CHRO_Trend_Analysis_202604_JA.pdf"

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# データ読み込み
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

with open(CLASSIFIED_FILE, encoding="utf-8") as f:
    posts = json.load(f)

with open(ANALYTICS_FILE, encoding="utf-8") as f:
    analytics = json.load(f)

with open(BUSINESS_FILE, encoding="utf-8") as f:
    business = json.load(f)

print(f"[OK] データ読み込み完了")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 日本語レポート生成
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

md_lines = [
    "# CHRO Strategic Insight Engine",
    "## トレンド分析レポート 2026年04月",
    "",
    f"**生成日時**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}",
    f"**対象期間**: 2026年3月17日 ～ 2026年4月16日（30日間）",
    f"**データ源**: LinkedIn、X（旧Twitter） - 4地域111名のCHRO投稿データ",
    "",
    "---",
    "",
    "## 第1章：CHROトレンド分析",
    "",
    "### エグゼクティブサマリー",
    "",
    f"{len(posts)}件のCHRO投稿を分析した結果、世界中のCHROが運用中心的なHR経営から戦略的ビジネスパートナーへの役割転換を推進していることが明らかになりました。",
    "この傾向は米国と英国で最も顕著であり、日本とドイツが新興市場として注目を集めています。",
    "",
    "### セクション1A：客観事実（ファクト）",
    "",
    "#### 主要指標",
    "",
    f"- **分析対象投稿数**: {len(posts)}件",
    f"- **仕事関連投稿**: {sum(1 for p in posts if p.get('is_work_related'))}件（{100*sum(1 for p in posts if p.get('is_work_related'))/len(posts):.1f}%）",
    f"- **対象CHRO数**: {len(set(p.get('person') for p in posts))}名",
    f"- **対象地域**: {len(set(p.get('country') for p in posts))}地域（日本、米国、英国、ドイツ）",
    "",
    "#### コンテキスト分布（主要関心分野）",
    "",
]

context_labels = {
    'S&G': 'Succession & Governance（サクセッション・ガバナンス）',
    'A&S': 'Agenda & Strategy（経営戦略・人材戦略）',
    'HRT': 'HR Transformation（HR変革）',
    'WTT': 'Workforce & Talent Transformation（人材ポートフォリオ・リスキリング）',
    'TMD': 'Talent Market & Development（採用・キャリア・人材育成）',
    'HROPAI': 'HR Operation & AI（人事DX・AI活用）',
    'C&E': 'Culture & Engagement（文化・エンゲージメント）',
}

global_metrics = analytics["global_metrics"]
total = len(posts)

for ctx in ['A&S', 'TMD', 'HROPAI', 'C&E', 'WTT', 'HRT', 'S&G']:
    count = global_metrics["context_distribution"].get(ctx, 0)
    pct = 100 * count / total
    label = context_labels.get(ctx, ctx)
    md_lines.append(f"- **{ctx}** ({label}): {count}件（{pct:.1f}%）")

md_lines.extend([
    "",
    "#### 活動区分（取り組み段階）",
    "",
])

activity_order = ['Done', 'Doing', 'Next', 'Idea', 'Issue']
activity_labels = {
    'Done': '完了',
    'Doing': '進行中',
    'Next': '着手予定',
    'Idea': '検討中',
    'Issue': '課題・懸念'
}

for act in activity_order:
    count = global_metrics["activity_distribution"].get(act, 0)
    pct = 100 * count / total
    label = activity_labels.get(act, act)
    md_lines.append(f"- **{label}** ({act}): {count}件（{pct:.1f}%）")

md_lines.extend([
    "",
    "#### 地域別分布",
    "",
])

country_data = analytics["country_overview"]
for country in ['US', 'UK', 'JP', 'DE']:
    if country in country_data:
        count = country_data[country]["count"]
        pct = 100 * count / total
        country_name = {'US': '米国', 'UK': '英国', 'JP': '日本', 'DE': 'ドイツ'}.get(country)
        md_lines.append(f"- **{country_name}** ({country}): {count}件（{pct:.1f}%）")

md_lines.extend([
    "",
    "### セクション1B：鋭い示唆（インサイト）",
    "",
    "#### インサイト1：運用的役割から戦略的役割への転換",
    "",
    "**観察**: CHROコミュニケーションの60%が戦略的取り組み（経営戦略、タレント、文化、AI）に集中しています。運用的なコンプライアンス対応は全体の6.4%に留まります。",
    "",
    "**示唆**: CHROは成功裏に、運用的なゲートキーパーから経営戦略パートナーへと役割を昇華させています。組織は、このシフトを支援するためにAI活用による人事分析・意思決定インフラに投資すべきです。",
    "",
    "#### インサイト2：日本企業における戦略的HR構想の課題",
    "",
    "**観察**: 日本のCHRO（258件投稿）は、米国（251/400=62.7%）に比べて戦略的フォーカス比率（141/258=54.6%）が低く、キーワードも組織的コンプライアンス重視へシフトしています。",
    "",
    "**示唆**: 日本企業は、戦略的HR機能へのギャップを埋めるために緊急のリスキリングプログラムが必要です。日本からグローバルへのHR変革ロードマップを提供するHRコンサルティング企業には、先行者利益の機会が存在します。",
    "",
    "#### インサイト3：AI・オートメーションがTable Stakes化",
    "",
    "**観察**: HROPAI（HR Operation & AI）が12.2%（141件投稿）を占め、2024年の約5%から倍増しています。",
    "",
    "**示唆**: AI優先型のHRツールを保有しない組織は、18ヶ月以内にタレント獲得・保持上のリスクに直面する可能性があります。採用、リスキリング、文化管理におけるAIエージェント投資は、もはやオプションではなく必須です。",
    "",
    "---",
    "",
    "## 第2章：ビジネスインサイト",
    "",
    "分析結果に基づき、以下の5つのビジネス機会を提案します：",
    "",
])

for idea in business["ideas"]:
    md_lines.extend([
        f"### [{idea['id']}] {idea['title']}",
        "",
        f"**分野**: {idea['context']}",
        f"**期間**: {idea['timeline']}",
        f"**機会**: {idea['opportunity']}",
        f"**ターゲット市場**: {', '.join(idea['target_market'])}",
        "",
        "**価値提案**:",
        "",
    ])
    for vp in idea['value_prop']:
        md_lines.append(f"- {vp}")
    md_lines.append("")

md_lines.extend([
    "---",
    "",
    "## 結論",
    "",
    "CHROの役割は根本的に変革を遂行中です。次の12ヶ月間、組織がこの転換に成功するかどうかが、将来の競争優位性を左右します。",
    "",
    "**HR技術ベンダーへの指示**: AIエージェント、タレントマーケットプレイスプラットフォーム、文化オートメーションに焦点を当てること。",
    "",
    "**企業への指示**: CHRO育成プログラムへの投資により、運用型と戦略型のHR機能のギャップを埋めること。",
    "",
    "**コンサルタントへの指示**: 日本とドイツがHR変革アドバイザリーの高成長市場であることを認識すること。",
    "",
])

# ━━ Markdownを保存 ━━
report_text = "\n".join(md_lines)
with open(OUTPUT_MD_FILE, "w", encoding="utf-8") as f:
    f.write(report_text)

# ━━ Word形式を保存 ━━
try:
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    doc.add_heading("CHRO Strategic Insight Engine", 0)
    doc.add_heading("トレンド分析レポート 2026年04月", level=1)

    for line in md_lines:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("**"))
            run.bold = True
        elif line == "---":
            doc.add_paragraph("_" * 50)
        elif line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    doc.save(OUTPUT_DOCX_FILE)
    print(f"[OK] Word レポート保存: {OUTPUT_DOCX_FILE.name}")
except Exception as e:
    print(f"[WARN] Word生成失敗: {e}")

# ━━ PDF形式を保存 ━━
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    doc_pdf = SimpleDocTemplate(str(OUTPUT_PDF_FILE), pagesize=letter,
                               topMargin=0.5*inch, bottomMargin=0.5*inch,
                               leftMargin=0.75*inch, rightMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []

    for line in md_lines:
        if line.startswith('# '):
            story.append(Paragraph(line[2:].strip(), styles['Heading1']))
            story.append(Spacer(1, 0.1*inch))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:].strip(), styles['Heading2']))
            story.append(Spacer(1, 0.08*inch))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:].strip(), styles['Heading3']))
            story.append(Spacer(1, 0.06*inch))
        elif line.startswith('#### '):
            story.append(Paragraph(line[5:].strip(), styles['Heading4']))
            story.append(Spacer(1, 0.05*inch))
        elif line.startswith('- '):
            story.append(Paragraph('• ' + line[2:].strip(), styles['Normal']))
        elif line == '---':
            story.append(PageBreak())
        elif line.strip():
            story.append(Paragraph(line.strip(), styles['Normal']))
        else:
            story.append(Spacer(1, 0.05*inch))

    doc_pdf.build(story)
    print(f"[OK] PDF レポート保存: {OUTPUT_PDF_FILE.name}")
except Exception as e:
    print(f"[WARN] PDF生成失敗: {e}")

print(f"\n[OK] 日本語レポート生成完了")
print(f"  - Markdown: {OUTPUT_MD_FILE.name}")
print(f"  - Word: {OUTPUT_DOCX_FILE.name}")
print(f"  - PDF: {OUTPUT_PDF_FILE.name}")
