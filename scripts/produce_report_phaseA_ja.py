#!/usr/bin/env python3
"""
produce_report_phaseA_ja.py
Phase A分析結果をWord/PDF形式で出力（日本語ビジュアルレポート）

出力:
  - analytics_phaseA_202604.docx (Word レポート)
  - analytics_phaseA_202604.pdf (PDF レポート)
"""
import json
import sys
from pathlib import Path
from datetime import datetime

# Windows UTF-8 対応
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DATA_DIR = Path(__file__).parent.parent / "data"
ANALYTICS_PHASEA_FILE = DATA_DIR / "analytics_phaseA_202604.json"
OUTPUT_DOCX_FILE = DATA_DIR / "analytics_phaseA_202604.docx"
OUTPUT_PDF_FILE = DATA_DIR / "analytics_phaseA_202604.pdf"

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# データ読み込み
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

with open(ANALYTICS_PHASEA_FILE, encoding="utf-8") as f:
    analytics = json.load(f)

print(f"[OK] Phase A データ読み込み完了")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Word形式で出力
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade_cell(cell, fill_color):
        """セルの背景色を設定"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def set_cell_border(cell, **kwargs):
        """セルのボーダーを設定"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            if edge in kwargs:
                edge_element = OxmlElement(f'w:{edge}')
                edge_element.set(qn('w:val'), kwargs[edge].get('style', 'single'))
                edge_element.set(qn('w:sz'), str(kwargs[edge].get('size', 4)))
                edge_element.set(qn('w:space'), str(kwargs[edge].get('space', 0)))
                edge_element.set(qn('w:color'), kwargs[edge].get('color', 'auto'))
                tcBorders.append(edge_element)
        tcPr.append(tcBorders)

    doc = Document()

    # タイトル
    title = doc.add_heading('CHRO Strategic Insight Engine', level=0)
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Phase A分析レポート', level=1)
    subtitle_format = subtitle.paragraph_format
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # メタデータ
    metadata = doc.add_paragraph()
    metadata.add_run('生成日時: ').bold = True
    metadata.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))

    period = doc.add_paragraph()
    period.add_run('対象期間: ').bold = True
    period.add_run('2026年3月17日～2026年4月16日（30日間）')

    doc.add_paragraph()  # 空行

    # ━━ Section 1: SNS情報サマリー ━━
    doc.add_heading('1. SNS情報サマリー（4企業群比較）', level=2)

    doc.add_paragraph('表1: CHRO人数・ポスト数・媒体別内訳', style='Heading 3')

    # SNS Summary Table
    sns_summary = analytics['sns_summary']
    table1 = doc.add_table(rows=5, cols=6)
    table1.style = 'Light Grid Accent 1'

    # ヘッダー
    header_cells = table1.rows[0].cells
    headers = ['企業群', 'CHRO人数', '総ポスト数', '仕事関連', 'Xポスト', 'LinkedInポスト']
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        shade_cell(header_cells[i], 'D5E8F0')
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # データ行
    for idx, country in enumerate(['JP', 'US', 'UK', 'DE']):
        row_cells = table1.rows[idx + 1].cells
        summary = sns_summary[country]
        row_data = [
            summary['country_ja'],
            str(summary['chro_count']),
            str(summary['total_posts']),
            str(summary['work_posts']),
            str(summary['x_posts']),
            str(summary['linkedin_posts'])
        ]
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
            for paragraph in row_cells[i].paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # インサイト
    doc.add_paragraph('地域別の特徴と洞察:', style='Heading 3')

    insights = [
        ('日本（JP）', 'LinkedInが大多数（96.5%）。プロフェッショナルネットワーク中心の発信。'),
        ('米国（US）', 'X と LinkedIn の両者が活発（X: 12.3%, LinkedIn: 87.8%）。言論の多元化が進んでいる。'),
        ('英国（UK）', 'LinkedIn優位（91.0%）。欧州ではLinkedIn文化が強い傾向。'),
        ('ドイツ（DE）', 'LinkedInが中心（88.1%）。独仏でも同傾向で、欧州の標準的な発信媒体。'),
    ]

    for region, insight_text in insights:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(region).bold = True
        p.add_run(f': {insight_text}')

    doc.add_page_break()

    # ━━ Section 2: コンテキスト分布（4テーブル）━━
    doc.add_heading('2. コンテキスト分布（4企業群比較）', level=2)

    context_labels = {
        'S&G': 'Succession & Governance',
        'A&S': 'Agenda & Strategy',
        'HRT': 'HR Transformation',
        'WTT': 'Workforce & Talent Transformation',
        'TMD': 'Talent Market & Development',
        'HROPAI': 'HR Operation & AI',
        'C&E': 'Culture & Engagement',
    }

    context_by_country = analytics['context_distribution']
    country_info = [
        ('JP', '日本', 'D5E8F0'),
        ('US', '米国', 'E8F0D5'),
        ('UK', '英国', 'F0D5E8'),
        ('DE', 'ドイツ', 'F0E8D5'),
    ]

    for country_code, country_name, header_color in country_info:
        doc.add_paragraph(f'表: {country_name}のコンテキスト別分布', style='Heading 3')

        # テーブル作成
        table = doc.add_table(rows=8, cols=4)
        table.style = 'Light Grid Accent 1'

        # ヘッダー
        header_cells = table.rows[0].cells
        header_texts = ['コンテキスト', 'ポスト数', '割合（%）', '説明']
        for i, header_text in enumerate(header_texts):
            header_cells[i].text = header_text
            shade_cell(header_cells[i], header_color)
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # データ行
        distribution = context_by_country[country_code]['distribution']
        ctx_order = ['A&S', 'TMD', 'HROPAI', 'C&E', 'WTT', 'HRT', 'S&G']

        for row_idx, ctx in enumerate(ctx_order):
            row_cells = table.rows[row_idx + 1].cells
            data = distribution[ctx]
            count = data['count']
            pct = data['percentage']

            row_cells[0].text = ctx
            row_cells[1].text = str(count)
            row_cells[2].text = f"{pct:.1f}%"
            row_cells[3].text = context_labels[ctx]

            # 中央寄せ
            for i in range(4):
                for paragraph in row_cells[i].paragraphs:
                    if i < 3:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        if country_code != 'DE':
            doc.add_page_break()

    doc.add_page_break()

    # ━━ Section 3: 4企業群の比較分析 ━━
    doc.add_heading('3. 4企業群の比較分析', level=2)

    doc.add_paragraph('コンテキスト別の地域比較', style='Heading 3')

    # A&S 比較
    doc.add_paragraph('A&S（経営戦略）における国別差異:', style='Heading 4')

    table_as = doc.add_table(rows=5, cols=2)
    table_as.style = 'Light Grid Accent 1'

    as_header = table_as.rows[0].cells
    as_header[0].text = '地域'
    as_header[1].text = 'A&S比率'
    for cell in as_header:
        shade_cell(cell, 'D5E8F0')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    as_data = [
        ('日本', '39.9%（戦略重視）'),
        ('米国', '41.5%（戦略最優先）'),
        ('英国', '43.5%（高い戦略意識）'),
        ('ドイツ', '43.5%（堅実な戦略志向）'),
    ]

    for idx, (country, pct_text) in enumerate(as_data):
        row = table_as.rows[idx + 1].cells
        row[0].text = country
        row[1].text = pct_text

    doc.add_paragraph()

    # HROPAI 比較
    doc.add_paragraph('HROPAI（HR DX・AI）における国別差異:', style='Heading 4')

    table_ai = doc.add_table(rows=5, cols=2)
    table_ai.style = 'Light Grid Accent 1'

    ai_header = table_ai.rows[0].cells
    ai_header[0].text = '地域'
    ai_header[1].text = 'HROPAI比率'
    for cell in ai_header:
        shade_cell(cell, 'E8F0D5')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    ai_data = [
        ('日本', '19.0%（急速成長中、テクノロジー志向が強い）'),
        ('米国', '10.5%（先進テクノロジー地）'),
        ('英国', '7.3%（DX構想の初期段階）'),
        ('ドイツ', '15.5%（デジタル化推進）'),
    ]

    for idx, (country, pct_text) in enumerate(ai_data):
        row = table_ai.rows[idx + 1].cells
        row[0].text = country
        row[1].text = pct_text

    doc.add_paragraph()
    doc.add_paragraph('【主要な発見】日本のCHROが HROPAI に最も高い関心を持っており、テクノロジー志向が強いことが判明しました。一方、経営戦略（A&S）比率は米英独に比べて若干低く、戦略的人事の深化に課題があることが示唆されています。')

    doc.add_page_break()

    # 次のフェーズ
    doc.add_heading('次のステップ', level=2)

    phases = [
        ('Phase B', 'コンテキスト×活動区分のマトリクス表（企業群ごとに4つ）'),
        ('Phase C', 'コンテキスト別のキーワードランキング + 深堀インサイト'),
    ]

    for phase_name, description in phases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(phase_name).bold = True
        p.add_run(f': {description}')

    doc.save(OUTPUT_DOCX_FILE)
    print(f"[OK] Word レポート保存: {OUTPUT_DOCX_FILE.name}")

except Exception as e:
    print(f"[WARN] Word生成失敗: {e}")
    import traceback
    traceback.print_exc()

# ━━ PDF形式を保存 ━━
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors

    pdf_file = SimpleDocTemplate(
        str(OUTPUT_PDF_FILE),
        pagesize=A4,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )

    styles = getSampleStyleSheet()
    story = []

    # タイトル
    story.append(Paragraph('CHRO Strategic Insight Engine', styles['Heading1']))
    story.append(Paragraph('Phase A分析レポート', styles['Heading2']))
    story.append(Spacer(1, 0.2*inch))

    # メタデータ
    story.append(Paragraph(f"<b>生成日時:</b> {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}", styles['Normal']))
    story.append(Paragraph('対象期間: 2026年3月17日～2026年4月16日（30日間）', styles['Normal']))
    story.append(Spacer(1, 0.2*inch))

    # SNS Summary
    story.append(Paragraph('1. SNS情報サマリー（4企業群比較）', styles['Heading2']))

    sns_summary = analytics['sns_summary']
    table_data = [['企業群', 'CHRO人数', '総ポスト数', '仕事関連', 'Xポスト', 'LinkedInポスト']]
    for country in ['JP', 'US', 'UK', 'DE']:
        summary = sns_summary[country]
        table_data.append([
            summary['country_ja'],
            str(summary['chro_count']),
            str(summary['total_posts']),
            str(summary['work_posts']),
            str(summary['x_posts']),
            str(summary['linkedin_posts'])
        ])

    table1 = Table(table_data)
    table1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(table1)
    story.append(Spacer(1, 0.3*inch))

    # Context Distribution
    story.append(Paragraph('2. コンテキスト分布（4企業群比較）', styles['Heading2']))

    context_by_country = analytics['context_distribution']
    countries = [('JP', '日本'), ('US', '米国'), ('UK', '英国'), ('DE', 'ドイツ')]

    for country_code, country_name in countries:
        story.append(Paragraph(f'{country_name}のコンテキスト別分布', styles['Heading3']))

        distribution = context_by_country[country_code]['distribution']
        ctx_order = ['A&S', 'TMD', 'HROPAI', 'C&E', 'WTT', 'HRT', 'S&G']

        table_data = [['コンテキスト', 'ポスト数', '割合（%）']]
        for ctx in ctx_order:
            data = distribution[ctx]
            table_data.append([ctx, str(data['count']), f"{data['percentage']:.1f}%"])

        table_ctx = Table(table_data)
        table_ctx.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(table_ctx)
        story.append(Spacer(1, 0.15*inch))

    story.append(PageBreak())

    # 比較分析
    story.append(Paragraph('3. 4企業群の比較分析', styles['Heading2']))

    story.append(Paragraph('A&S（経営戦略）における国別差異:', styles['Heading3']))
    story.append(Paragraph('• 日本: 39.9%（戦略重視）', styles['Normal']))
    story.append(Paragraph('• 米国: 41.5%（戦略最優先）', styles['Normal']))
    story.append(Paragraph('• 英国: 43.5%（高い戦略意識）', styles['Normal']))
    story.append(Paragraph('• ドイツ: 43.5%（堅実な戦略志向）', styles['Normal']))
    story.append(Spacer(1, 0.15*inch))

    story.append(Paragraph('HROPAI（HR DX・AI）における国別差異:', styles['Heading3']))
    story.append(Paragraph('• 日本: 19.0%（急速成長中、テクノロジー志向が強い）', styles['Normal']))
    story.append(Paragraph('• 米国: 10.5%（先進テクノロジー地）', styles['Normal']))
    story.append(Paragraph('• 英国: 7.3%（DX構想の初期段階）', styles['Normal']))
    story.append(Paragraph('• ドイツ: 15.5%（デジタル化推進）', styles['Normal']))
    story.append(Spacer(1, 0.2*inch))

    story.append(Paragraph('【主要な発見】日本のCHROが HROPAI に最も高い関心を持っており、テクノロジー志向が強いことが判明しました。一方、経営戦略（A&S）比率は米英独に比べて若干低く、戦略的人事の深化に課題があることが示唆されています。', styles['Normal']))

    pdf_file.build(story)
    print(f"[OK] PDF レポート保存: {OUTPUT_PDF_FILE.name}")

except Exception as e:
    print(f"[WARN] PDF生成失敗: {e}")
    import traceback
    traceback.print_exc()

print(f"\n[OK] Phase A ビジュアルレポート生成完了")
print(f"  - Word: {OUTPUT_DOCX_FILE.name}")
print(f"  - PDF: {OUTPUT_PDF_FILE.name}")
