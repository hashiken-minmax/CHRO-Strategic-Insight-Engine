#!/usr/bin/env python3
"""
produce_report_phaseA_with_charts.py
Phase A分析結果をWord形式で出力（円グラフ - 扇形塗りつぶし）

出力:
  - analytics_phaseA_202604.docx (Word レポート - 4つの円グラフ)
"""
import json
import sys
from pathlib import Path
from datetime import datetime
import io
import math

# Windows UTF-8 対応
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DATA_DIR = Path(__file__).parent.parent / "data"
ANALYTICS_PHASEA_FILE = DATA_DIR / "analytics_phaseA_202604.json"
OUTPUT_DOCX_FILE = DATA_DIR / "analytics_phaseA_202604.docx"
TEMP_DIR = DATA_DIR / "temp_charts"
TEMP_DIR.mkdir(exist_ok=True)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# データ読み込み
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

with open(ANALYTICS_PHASEA_FILE, encoding="utf-8") as f:
    analytics = json.load(f)

print(f"[OK] Phase A データ読み込み完了")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PIL で円グラフ生成（扇形塗りつぶし）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

from PIL import Image, ImageDraw, ImageFont

context_by_country = analytics['context_distribution']
countries = [('JP', '日本'), ('US', '米国'), ('UK', '英国'), ('DE', 'ドイツ')]

# 色パレット（見やすい配色）
colors_list = [
    '#FF6B6B', '#4ECDC4', '#45B7D1', '#FFA07A',
    '#98D8C8', '#F7DC6F', '#BB8FCE', '#85C1E2'
]

def hex_to_rgb(hex_color):
    """HexカラーをRGBに変換"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def create_pie_chart_image(country_code, country_name, distribution, output_file):
    """PIL を使用して扇形塗りつぶされた円グラフを作成"""
    try:
        # 画像作成（背景白）
        img = Image.new('RGB', (480, 520), color='white')
        draw = ImageDraw.Draw(img)

        ctx_order = ['A&S', 'TMD', 'HROPAI', 'C&E', 'WTT', 'HRT', 'S&G']

        labels_and_sizes = []
        for ctx in ctx_order:
            data = distribution[ctx]
            count = data['count']
            if count > 0:
                labels_and_sizes.append((ctx, count, data['percentage']))

        # 円グラフのパラメータ
        center_x, center_y = 240, 200
        radius = 140

        # タイトル描画
        title_text = f'{country_name}（{country_code}）'
        try:
            # フォントサイズを指定
            font = ImageFont.load_default()
        except:
            font = ImageFont.load_default()

        title_bbox = draw.textbbox((0, 0), title_text, font=font)
        title_width = title_bbox[2] - title_bbox[0]
        draw.text(((480 - title_width) // 2, 15), title_text, fill='black', font=font)

        # 扇形を描画（塗りつぶし）
        total = sum(size for _, size, _ in labels_and_sizes)

        start_angle = -90  # 上から開始
        for idx, (label, size, percentage) in enumerate(labels_and_sizes):
            angle = (size / total) * 360
            end_angle = start_angle + angle

            color = colors_list[idx % len(colors_list)]
            color_rgb = hex_to_rgb(color)

            # 扇形のポイントリストを作成
            points = [(center_x, center_y)]

            # 円弧上のポイントを計算
            steps = max(int(angle / 2), 2)  # 滑らかさを調整
            for step in range(steps + 1):
                current_angle = start_angle + (angle * step / steps)
                rad = math.radians(current_angle)
                x = center_x + radius * math.cos(rad)
                y = center_y + radius * math.sin(rad)
                points.append((int(x), int(y)))

            # 扇形を塗りつぶし
            if len(points) > 2:
                draw.polygon(points, fill=color_rgb, outline='white')

            # パーセンテージをテキスト表示（扇形の中央）
            mid_angle = start_angle + angle / 2
            rad = math.radians(mid_angle)
            text_radius = radius * 0.65
            text_x = center_x + text_radius * math.cos(rad)
            text_y = center_y + text_radius * math.sin(rad)

            pct_text = f'{percentage:.0f}%'
            pct_bbox = draw.textbbox((0, 0), pct_text, font=font)
            pct_width = pct_bbox[2] - pct_bbox[0]
            pct_height = pct_bbox[3] - pct_bbox[1]

            draw.text(
                (int(text_x - pct_width / 2), int(text_y - pct_height / 2)),
                pct_text,
                fill='white',
                font=font
            )

            start_angle = end_angle

        # 凡例を描画
        legend_y = 360
        for idx, (label, size, percentage) in enumerate(labels_and_sizes):
            color = colors_list[idx % len(colors_list)]
            color_rgb = hex_to_rgb(color)

            # 色付きボックス
            legend_box_x = 30
            legend_box_y = legend_y + idx * 22
            draw.rectangle(
                [legend_box_x, legend_box_y, legend_box_x + 12, legend_box_y + 12],
                fill=color_rgb,
                outline='black'
            )

            # テキスト
            legend_text = f'{label} ({size}件, {percentage:.1f}%)'
            draw.text(
                (legend_box_x + 18, legend_box_y),
                legend_text,
                fill='black',
                font=font
            )

        img.save(str(output_file), 'PNG')
        return True

    except Exception as e:
        print(f"[ERR] 円グラフ作成エラー: {e}")
        import traceback
        traceback.print_exc()
        return False


chart_files = {}

for country_code, country_name in countries:
    distribution = context_by_country[country_code]['distribution']

    chart_file = TEMP_DIR / f'chart_{country_code}.png'

    if create_pie_chart_image(country_code, country_name, distribution, chart_file):
        chart_files[country_code] = chart_file
        print(f"[OK] 円グラフ生成: {country_code}")
    else:
        print(f"[WARN] 円グラフ生成失敗: {country_code}")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Word形式で出力
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade_cell(cell, fill_color):
        """セルの背景色を設定"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    doc = Document()

    # タイトル
    title = doc.add_heading('CHRO Strategic Insight Engine', level=0)
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Phase A分析レポート', level=1)
    subtitle_format = subtitle.paragraph_format
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # メタデータ
    metadata = doc.add_paragraph()
    metadata.add_run('生成日時: ').bold = True
    metadata.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))

    period = doc.add_paragraph()
    period.add_run('対象期間: ').bold = True
    period.add_run('2026年3月17日～2026年4月16日（30日間）')

    doc.add_paragraph()  # 空行

    # ━━ Section 1: SNS情報サマリー ━━
    doc.add_heading('1. SNS情報サマリー（4企業群比較）', level=2)

    doc.add_paragraph('表1: CHRO人数・ポスト数・媒体別内訳', style='Heading 3')

    # SNS Summary Table
    sns_summary = analytics['sns_summary']
    table1 = doc.add_table(rows=5, cols=6)
    table1.style = 'Light Grid Accent 1'

    # ヘッダー
    header_cells = table1.rows[0].cells
    headers = ['企業群', 'CHRO人数', '総ポスト数', '仕事関連', 'Xポスト', 'LinkedInポスト']
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        shade_cell(header_cells[i], 'D5E8F0')
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # データ行
    for idx, country in enumerate(['JP', 'US', 'UK', 'DE']):
        row_cells = table1.rows[idx + 1].cells
        summary = sns_summary[country]
        row_data = [
            summary['country_ja'],
            str(summary['chro_count']),
            str(summary['total_posts']),
            str(summary['work_posts']),
            str(summary['x_posts']),
            str(summary['linkedin_posts'])
        ]
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
            for paragraph in row_cells[i].paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # インサイト
    doc.add_paragraph('地域別の特徴と洞察:', style='Heading 3')

    insights = [
        ('日本（JP）', 'LinkedInが大多数（96.5%）。プロフェッショナルネットワーク中心の発信。'),
        ('米国（US）', 'X と LinkedIn の両者が活発（X: 12.3%, LinkedIn: 87.8%）。言論の多元化が進んでいる。'),
        ('英国（UK）', 'LinkedIn優位（91.0%）。欧州ではLinkedIn文化が強い傾向。'),
        ('ドイツ（DE）', 'LinkedInが中心（88.1%）。独仏でも同傾向で、欧州の標準的な発信媒体。'),
    ]

    for region, insight_text in insights:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(region).bold = True
        p.add_run(f': {insight_text}')

    doc.add_page_break()

    # ━━ Section 2: コンテキスト分布（円グラフ）━━
    doc.add_heading('2. コンテキスト分布（円グラフ - 4企業群比較）', level=2)

    doc.add_paragraph('各国のCHRO投稿が示すコンテキスト別の関心分野を円グラフで可視化。1ページに4つの円グラフを配置。', style='Normal')
    doc.add_paragraph()

    # 2x2 レイアウト: テーブル内に4つの円グラフを配置
    chart_table = doc.add_table(rows=2, cols=2)
    chart_table.autofit = False

    # 各セルに円グラフを挿入
    chart_positions = [
        (0, 0, 'JP'),
        (0, 1, 'US'),
        (1, 0, 'UK'),
        (1, 1, 'DE')
    ]

    for row_idx, col_idx, country_code in chart_positions:
        cell = chart_table.rows[row_idx].cells[col_idx]

        # セルのクリア
        cell.text = ''

        # 画像を挿入
        if country_code in chart_files:
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(str(chart_files[country_code]), width=Inches(3.2))
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ━━ Section 3: 数値比較分析 ━━
    doc.add_heading('3. 4企業群の比較分析', level=2)

    doc.add_paragraph('コンテキスト別の地域比較', style='Heading 3')

    # A&S 比較
    doc.add_paragraph('A&S（経営戦略）における国別差異:', style='Heading 4')

    table_as = doc.add_table(rows=5, cols=2)
    table_as.style = 'Light Grid Accent 1'

    as_header = table_as.rows[0].cells
    as_header[0].text = '地域'
    as_header[1].text = 'A&S比率'
    for cell in as_header:
        shade_cell(cell, 'D5E8F0')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    as_data = [
        ('日本', '39.9%（戦略重視）'),
        ('米国', '41.5%（戦略最優先）'),
        ('英国', '43.5%（高い戦略意識）'),
        ('ドイツ', '43.5%（堅実な戦略志向）'),
    ]

    for idx, (country, pct_text) in enumerate(as_data):
        row = table_as.rows[idx + 1].cells
        row[0].text = country
        row[1].text = pct_text

    doc.add_paragraph()

    # HROPAI 比較
    doc.add_paragraph('HROPAI（HR DX・AI）における国別差異:', style='Heading 4')

    table_ai = doc.add_table(rows=5, cols=2)
    table_ai.style = 'Light Grid Accent 1'

    ai_header = table_ai.rows[0].cells
    ai_header[0].text = '地域'
    ai_header[1].text = 'HROPAI比率'
    for cell in ai_header:
        shade_cell(cell, 'E8F0D5')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    ai_data = [
        ('日本', '19.0%（急速成長中、テクノロジー志向が強い）'),
        ('米国', '10.5%（先進テクノロジー地）'),
        ('英国', '7.3%（DX構想の初期段階）'),
        ('ドイツ', '15.5%（デジタル化推進）'),
    ]

    for idx, (country, pct_text) in enumerate(ai_data):
        row = table_ai.rows[idx + 1].cells
        row[0].text = country
        row[1].text = pct_text

    doc.add_paragraph()

    insight_para = doc.add_paragraph()
    insight_para.add_run('【主要な発見】').bold = True
    insight_para.add_run('日本のCHROが HROPAI に最も高い関心を持っており、テクノロジー志向が強いことが判明しました。一方、経営戦略（A&S）比率は米英独に比べて若干低く、戦略的人事の深化に課題があることが示唆されています。')

    doc.add_page_break()

    # 次のフェーズ
    doc.add_heading('次のステップ', level=2)

    phases = [
        ('Phase B', 'コンテキスト×活動区分のマトリクス表（企業群ごとに4つ）'),
        ('Phase C', 'コンテキスト別のキーワードランキング + 深堀インサイト'),
    ]

    for phase_name, description in phases:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(phase_name).bold = True
        p.add_run(f': {description}')

    doc.save(OUTPUT_DOCX_FILE)
    print(f"[OK] Word レポート保存: {OUTPUT_DOCX_FILE.name}")
    print(f"     （扇形塗りつぶし円グラフ4つ，1ページ配置）")

except Exception as e:
    print(f"[WARN] Word生成失敗: {e}")
    import traceback
    traceback.print_exc()

print(f"\n[OK] Phase A ビジュアルレポート（円グラフ版）生成完了")
