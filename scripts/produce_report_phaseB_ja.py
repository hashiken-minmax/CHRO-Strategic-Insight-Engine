#!/usr/bin/env python3
"""
produce_report_phaseB_ja.py
Phase B分析結果をWord形式で出力（マトリクス表）

出力:
  - analytics_phaseB_202604.docx (Word レポート - 4つのマトリクス表)
"""
import json
import sys
from pathlib import Path
from datetime import datetime

# Windows UTF-8 対応
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DATA_DIR = Path(__file__).parent.parent / "data"
ANALYTICS_PHASEB_FILE = DATA_DIR / "analytics_phaseB_202604.json"
OUTPUT_DOCX_FILE = DATA_DIR / "analytics_phaseB_202604.docx"

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# データ読み込み
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

with open(ANALYTICS_PHASEB_FILE, encoding="utf-8") as f:
    analytics = json.load(f)

print(f"[OK] Phase B データ読み込み完了")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Word形式で出力
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade_cell(cell, fill_color):
        """セルの背景色を設定"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def set_cell_text_alignment(cell, alignment):
        """セル内のテキスト配置を設定"""
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.alignment = alignment

    doc = Document()

    # タイトル
    title = doc.add_heading('CHRO Strategic Insight Engine', level=0)
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Phase B分析レポート', level=1)
    subtitle_format = subtitle.paragraph_format
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # メタデータ
    metadata = doc.add_paragraph()
    metadata.add_run('生成日時: ').bold = True
    metadata.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))

    period = doc.add_paragraph()
    period.add_run('対象期間: ').bold = True
    period.add_run('2026年3月17日～2026年4月16日（30日間）')

    doc.add_paragraph()

    # ━━ タイトル ━━
    doc.add_heading('コンテキスト × 活動区分のマトリクス分析', level=2)

    doc.add_paragraph('各国のCHRO投稿を「コンテキスト」（7分類）と「活動区分」（5段階）の2軸でクロス集計。各企業群の戦略的フォーカスと実行段階を可視化。')
    doc.add_paragraph()

    # 定義
    doc.add_heading('用語定義', level=3)
    p = doc.add_paragraph('【コンテキスト】', style='List Bullet')
    p.add_run('A&S（経営戦略）、TMD（採用・育成）、HROPAI（HR DX・AI）、C&E（文化・エンゲージメント）、WTT（人材ポートフォリオ）、HRT（HR変革）、S&G（サクセッション・ガバナンス）')

    p = doc.add_paragraph('【活動区分】', style='List Bullet')
    p.add_run('Done（完了）、Doing（進行中）、Next（着手予定）、Idea（検討中）、Issue（課題・懸念）')
    doc.add_paragraph()

    # ━━ 4つのマトリクス表 ━━
    matrix_by_country = analytics['matrix_by_country']
    country_names = {'JP': '日本', 'US': '米国', 'UK': '英国', 'DE': 'ドイツ'}
    ctx_order = ['A&S', 'TMD', 'HROPAI', 'C&E', 'WTT', 'HRT', 'S&G']
    activity_order = ['Done', 'Doing', 'Next', 'Idea', 'Issue']

    for idx, country in enumerate(['JP', 'US', 'UK', 'DE']):
        data = matrix_by_country[country]
        matrix = data['matrix']
        row_totals = data['row_totals']
        col_totals = data['col_totals']
        grand_total = data['grand_total']

        country_ja = country_names[country]

        # サブタイトル
        doc.add_heading(f'{country_ja}（{country}）のマトリクス', level=3)
        doc.add_paragraph(f'総ポスト数: {grand_total}件', style='Normal')

        # テーブル作成（8行×7列：ヘッダー+7コンテキスト、列：コンテキスト+5活動区分+計）
        table = doc.add_table(rows=9, cols=7)
        table.style = 'Light Grid Accent 1'

        # ヘッダー行
        header_row = table.rows[0]
        header_cells = header_row.cells

        headers = ['コンテキスト', 'Done', 'Doing', 'Next', 'Idea', 'Issue', '計']
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            shade_cell(header_cells[i], 'D5E8F0')
            set_cell_text_alignment(header_cells[i], WD_ALIGN_PARAGRAPH.CENTER)
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # データ行
        for row_idx, ctx in enumerate(ctx_order):
            row = table.rows[row_idx + 1]
            cells = row.cells

            # コンテキスト名
            cells[0].text = ctx
            set_cell_text_alignment(cells[0], WD_ALIGN_PARAGRAPH.CENTER)

            # 各活動区分のセル
            for col_idx, act in enumerate(activity_order):
                count = matrix[ctx][act]
                cells[col_idx + 1].text = str(count)
                set_cell_text_alignment(cells[col_idx + 1], WD_ALIGN_PARAGRAPH.CENTER)

                # セルの背景色（数値に応じてグラデーション）
                if count > 0:
                    # 色のグラデーション（0-最大値）
                    max_in_col = max(matrix[c][act] for c in ctx_order)
                    intensity = int(200 * count / max_in_col)
                    color = f'{220:02x}{int(220-intensity):02x}{int(220-intensity):02x}'
                    shade_cell(cells[col_idx + 1], color)

            # 計
            total = row_totals[ctx]
            cells[6].text = str(total)
            set_cell_text_alignment(cells[6], WD_ALIGN_PARAGRAPH.CENTER)
            shade_cell(cells[6], 'E8F0D5')
            for paragraph in cells[6].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # 合計行
        total_row = table.rows[8]
        total_cells = total_row.cells
        total_cells[0].text = '合計'
        shade_cell(total_cells[0], 'E8F0D5')
        set_cell_text_alignment(total_cells[0], WD_ALIGN_PARAGRAPH.CENTER)

        for col_idx, act in enumerate(activity_order):
            total_cells[col_idx + 1].text = str(col_totals[act])
            shade_cell(total_cells[col_idx + 1], 'E8F0D5')
            set_cell_text_alignment(total_cells[col_idx + 1], WD_ALIGN_PARAGRAPH.CENTER)
            for paragraph in total_cells[col_idx + 1].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        total_cells[6].text = str(grand_total)
        shade_cell(total_cells[6], 'FFF2CC')
        set_cell_text_alignment(total_cells[6], WD_ALIGN_PARAGRAPH.CENTER)
        for paragraph in total_cells[6].paragraphs:
            for run in paragraph.runs:
                run.bold = True

        # インサイト
        max_ctx = max(ctx_order, key=lambda x: row_totals[x])
        max_ctx_count = row_totals[max_ctx]
        max_act = max(activity_order, key=lambda x: col_totals[x])
        max_act_count = col_totals[max_act]

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('【主要指標】').bold = True
        doc.add_paragraph(f'最多コンテキスト: {max_ctx} ({max_ctx_count}件, {100*max_ctx_count/grand_total:.1f}%)', style='List Bullet')
        doc.add_paragraph(f'最多活動区分: {max_act} ({max_act_count}件, {100*max_act_count/grand_total:.1f}%)', style='List Bullet')

        if idx < 3:
            doc.add_page_break()

    doc.add_page_break()

    # ━━ 全体比較分析 ━━
    doc.add_heading('4企業群の比較分析', level=2)

    # コンテキスト別比較
    doc.add_heading('各コンテキストの国別比較', level=3)

    comp_table = doc.add_table(rows=8, cols=5)
    comp_table.style = 'Light Grid Accent 1'

    # ヘッダー
    header_cells = comp_table.rows[0].cells
    comp_headers = ['コンテキスト', '日本', '米国', '英国', 'ドイツ']
    for i, header_text in enumerate(comp_headers):
        header_cells[i].text = header_text
        shade_cell(header_cells[i], 'D5E8F0')
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # コンテキスト行
    for row_idx, ctx in enumerate(ctx_order):
        row = comp_table.rows[row_idx + 1]
        cells = row.cells

        cells[0].text = ctx
        for col_idx, country in enumerate(['JP', 'US', 'UK', 'DE']):
            count = matrix_by_country[country]['row_totals'][ctx]
            total = matrix_by_country[country]['grand_total']
            pct = 100 * count / total if total > 0 else 0
            cells[col_idx + 1].text = f'{count}\n({pct:.1f}%)'
            set_cell_text_alignment(cells[col_idx + 1], WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # 次のフェーズ
    doc.add_heading('次のステップ', level=2)

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Phase C: ').bold = True
    p.add_run('コンテキスト別のキーワードランキング + 深堀インサイト')

    doc.save(OUTPUT_DOCX_FILE)
    print(f"[OK] Word レポート保存: {OUTPUT_DOCX_FILE.name}")
    print(f"     （マトリクス表4つ + 比較分析）")

except Exception as e:
    print(f"[WARN] Word生成失敗: {e}")
    import traceback
    traceback.print_exc()

print(f"\n[OK] Phase B ビジュアルレポート生成完了")
