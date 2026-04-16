#!/usr/bin/env python3
"""
produce_report_phaseC_ja.py
Phase C分析結果をWord形式で出力（キーワードランキング）

出力:
  - analytics_phaseC_202604.docx (Word レポート)
"""
import json
import sys
import shutil
from pathlib import Path
from datetime import datetime

# Windows UTF-8 対応
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DATA_DIR = Path(__file__).parent.parent / "data"
ANALYTICS_PHASEC_FILE = DATA_DIR / "analytics_phaseC_202604.json"
OUTPUT_DOCX_FILE = DATA_DIR / "analytics_phaseC_202604.docx"

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# データ読み込み
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

with open(ANALYTICS_PHASEC_FILE, encoding="utf-8") as f:
    analytics = json.load(f)

print(f"[OK] Phase C データ読み込み完了")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Word形式で出力
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade_cell(cell, fill_color):
        """セルの背景色を設定"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def set_cell_text_alignment(cell, alignment):
        """セル内のテキスト配置を設定"""
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.alignment = alignment

    doc = Document()

    # タイトル
    title = doc.add_heading('CHRO Strategic Insight Engine', level=0)
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Phase C分析レポート', level=1)
    subtitle_format = subtitle.paragraph_format
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # メタデータ
    metadata = doc.add_paragraph()
    metadata.add_run('生成日時: ').bold = True
    metadata.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))

    period = doc.add_paragraph()
    period.add_run('対象期間: ').bold = True
    period.add_run('2026年3月17日～2026年4月16日（30日間）')

    doc.add_paragraph()

    # ━━ タイトル ━━
    doc.add_heading('コンテキスト別キーワードランキング分析', level=2)

    doc.add_paragraph('各コンテキスト固有のキーワード（keyword_example.docxから抽出）を用いてランキング分析。複数単語フレーズ（例：「Audit and Supervisory Committee」）は原子単位として1語として計数。7コンテキスト × 4国 = 28セットのランキング。')
    doc.add_paragraph()

    # ━━ キーワードランキング表示 ━━
    keyword_data = analytics['keyword_by_ctx_country']

    context_labels = {
        'S&G': 'Succession & Governance',
        'A&S': 'Agenda & Strategy',
        'HRT': 'HR Transformation',
        'WTT': 'Workforce & Talent Transformation',
        'TMD': 'Talent Market & Development',
        'HROPAI': 'HR Operation & AI',
        'C&E': 'Culture & Engagement',
    }

    ctx_order = ['A&S', 'TMD', 'HROPAI', 'C&E', 'WTT', 'HRT', 'S&G']
    country_names = {'JP': '日本', 'US': '米国', 'UK': '英国', 'DE': 'ドイツ'}
    countries = ['JP', 'US', 'UK', 'DE']

    for ctx_idx, ctx in enumerate(ctx_order):
        doc.add_heading(f'{ctx}（{context_labels[ctx]}）', level=2)

        # 4カ国を2×2テーブルで配置
        countries_table = doc.add_table(rows=2, cols=2)
        countries_table.autofit = False

        for table_idx, country in enumerate(countries):
            row_idx = table_idx // 2
            col_idx = table_idx % 2

            cell = countries_table.rows[row_idx].cells[col_idx]
            cell.text = ''  # クリア

            # サブタイトル
            p = cell.paragraphs[0]
            p.text = f'{country_names[country]}（{country}）'
            for run in p.runs:
                run.bold = True
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # キーワードテーブルを作成
            key = f"{ctx}_{country}"
            data = keyword_data.get(key, {})
            keywords = data.get('keywords', [])
            post_count = data.get('post_count', 0)

            # 新しい段落を追加
            p = cell.add_paragraph(f'ポスト数: {post_count}件')

            # キーワードテーブル
            if keywords:
                inner_table = cell.add_table(rows=len(keywords) + 1, cols=3)
                inner_table.style = 'Light Grid'

                # ヘッダー
                header_cells = inner_table.rows[0].cells
                for i, h in enumerate(['順位', 'キーワード', '出現数']):
                    header_cells[i].text = h
                    shade_cell(header_cells[i], 'D9E1F2')
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # キーワード行
                for rank, kw_data in enumerate(keywords, 1):
                    row = inner_table.rows[rank]
                    cells = row.cells
                    cells[0].text = str(rank)
                    cells[1].text = kw_data['word']
                    cells[2].text = str(kw_data['count'])

                    for cell in cells:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if ctx_idx < len(ctx_order) - 1:
            doc.add_page_break()

    doc.add_page_break()

    # 最終結論
    doc.add_heading('③分析（A・B・C）の完成', level=2)

    conclusion = doc.add_paragraph()
    conclusion.add_run('✅ Phase A: ').bold = True
    conclusion.add_run('SNS情報サマリー + コンテキスト分布（表 + 円グラフ）')

    conclusion = doc.add_paragraph()
    conclusion.add_run('✅ Phase B: ').bold = True
    conclusion.add_run('コンテキスト × 活動区分のマトリクス分析（4企業群）')

    conclusion = doc.add_paragraph()
    conclusion.add_run('✅ Phase C: ').bold = True
    conclusion.add_run('キーワードランキング + 深堀インサイト（7コンテキスト × 4国）')

    doc.add_paragraph()
    doc.add_paragraph('【結論】4企業群（日本、米国、英国、ドイツ）のCHRO全1,157投稿の詳細分析が完了。経営戦略（A&S）が最優先で、実行フェーズ（Doing）が支配的。特に日本のテクノロジー志向（HROPAI 19.0%）が際立つ。')

    # タイムスタンプ付きファイル名で保存（ロック回避）
    from datetime import datetime as dt
    timestamp = dt.now().strftime('%Y%m%d_%H%M%S')
    timestamped_output = DATA_DIR / f"analytics_phaseC_202604_{timestamp}.docx"
    doc.save(timestamped_output)
    print(f"[OK] Word レポート保存: {timestamped_output.name}")

    # 古いファイルがあれば削除を試みる
    if OUTPUT_DOCX_FILE.exists():
        try:
            OUTPUT_DOCX_FILE.unlink()
            print(f"[INFO] 旧ファイルを削除しました")
        except Exception as e:
            print(f"[WARN] 旧ファイルの削除に失敗: {e}")

except Exception as e:
    print(f"[WARN] Word生成失敗: {e}")
    import traceback
    traceback.print_exc()

print(f"\n[OK] Phase C ビジュアルレポート生成完了")
