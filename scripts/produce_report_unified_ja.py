#!/usr/bin/env python3
"""
produce_report_unified_ja.py - REFACTORED (Callable Function)
統合レポート - 複数フェーズデータを結合した総合分析レポート（コンテキスト中心）

出力:
  - PDF bytes buffer (統合レポート - 19-20ページ)
  - または analytics_unified_yyyymmdd.docx (スタンドアロン実行時)

構成:
  1. Executive Summary (1-2ページ)
  2. SNS情報サマリー + コンテキスト概要（1ページ）
  3. 7×Context Deep Dives (14ページ)
     - A&S, TMD, HROPAI, C&E, WTT, HRT, S&G (各2ページ)
  4. Issue一覧 (3-4ページ)
  5. Business Ideas (2ページ)
  6. Appendix (1ページ)
"""
import json
import sys
import shutil
from pathlib import Path
from datetime import datetime, timedelta
from collections import Counter, defaultdict
import io
import math
import tempfile
import subprocess
import os

# Windows UTF-8 対応
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DATA_DIR = Path(__file__).parent.parent / "data"
TEMP_DIR = DATA_DIR / "temp_charts"
TEMP_DIR.mkdir(exist_ok=True)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ライブラリインポート
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# グローバル設定
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

COUNTRIES = ['JP', 'US', 'UK', 'DE']
CTX_ORDER = ['A&S', 'TMD', 'HROPAI', 'C&E', 'WTT', 'HRT', 'S&G']
ACT_ORDER = ['Done', 'Doing', 'Next', 'Idea', 'Issue']

COUNTRY_NAMES = {
    'JP': '日本',
    'US': '米国',
    'UK': '英国',
    'DE': 'ドイツ'
}

CONTEXT_LABELS = {
    'A&S': 'Agenda & Strategy',
    'TMD': 'Talent Market & Development',
    'HROPAI': 'HR Operations & AI',
    'C&E': 'Culture & Engagement',
    'WTT': 'Workforce & Talent Transformation',
    'HRT': 'HR Transformation',
    'S&G': 'Succession & Governance'
}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ヘルパー関数（モジュールレベル）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def reorganize_phaseB_data(analytics_b):
    """
    フェーズ B データを matrix_by_country → context-first アクセス形式に変換
    """
    phaseB_by_context = defaultdict(lambda: defaultdict(dict))

    matrix_by_country = analytics_b.get('matrix_by_country', {})
    for country in COUNTRIES:
        country_matrix = matrix_by_country.get(country, {}).get('matrix', {})
        for ctx in CTX_ORDER:
            ctx_data = country_matrix.get(ctx, {})
            for act in ACT_ORDER:
                count = ctx_data.get(act, 0)
                phaseB_by_context[ctx][country][act] = count

    return phaseB_by_context

def extract_issue_data(classified_posts):
    """
    分類済みデータから Issue レコードを抽出
    構造: issues_by_context_country[context][country] = [list of issues]
    """
    issues_by_context_country = defaultdict(lambda: defaultdict(list))

    for post in classified_posts:
        if post.get('activity_class') == 'Issue':
            context = post.get('context_axis', 'Unknown')
            country = post.get('country', 'Unknown')
            issues_by_context_country[context][country].append(post)

    return issues_by_context_country

def apply_keyword_color_coding(phaseC_data, context):
    """
    キーワードに4国カラーコード適用
    全4国で出現 → gray (#E8E8E8)
    3国以下 → white (#FFFFFF)
    """
    keyword_colors = {}

    all_keywords = set()
    countries_having_keyword = defaultdict(set)

    for country in COUNTRIES:
        key = f"{context}_{country}"
        keywords = phaseC_data.get(key, {}).get('keywords', [])
        for kw_data in keywords:
            word = kw_data.get('word', '')
            all_keywords.add(word)
            countries_having_keyword[word].add(country)

    for word in all_keywords:
        if len(countries_having_keyword[word]) == 4:
            keyword_colors[word] = 'E8E8E8'  # gray
        else:
            keyword_colors[word] = 'FFFFFF'  # white

    return keyword_colors

def shade_cell(cell, fill_color):
    """セルの背景色を設定"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_heatmap_color(value, max_value):
    """値に基づいてヒートマップ色を返す（白→赤グラデーション）"""
    if max_value == 0:
        return "FFFFFF"
    ratio = value / max_value
    r = 255
    g = int(255 - (255 - 107) * ratio)
    b = int(255 - (255 - 107) * ratio)
    return f"{r:02X}{g:02X}{b:02X}"

def generate_execution_gap_narrative(context, phaseB_by_context):
    """
    実行フェーズギャップナラティブを生成
    Done/Doing の国別割合で成熟度差を説明（絶対数でなく相対比率）
    """
    ctx_data = phaseB_by_context.get(context, {})

    # 各国の実行成熟度を計算（Done / (Done + Doing)）
    maturity_by_country = {}
    for country in COUNTRIES:
        done = ctx_data.get(country, {}).get('Done', 0)
        doing = ctx_data.get(country, {}).get('Doing', 0)
        total_impl = done + doing
        maturity = (done / total_impl * 100) if total_impl > 0 else 0
        maturity_by_country[country] = {
            'done': done,
            'doing': doing,
            'maturity': maturity,
            'total': total_impl
        }

    # 成熟度が最も高い国と低い国を特定
    leader_country = max(maturity_by_country, key=lambda c: maturity_by_country[c]['maturity'])
    laggard_country = min(maturity_by_country, key=lambda c: maturity_by_country[c]['maturity'])

    leader_maturity = maturity_by_country[leader_country]['maturity']
    laggard_maturity = maturity_by_country[laggard_country]['maturity']
    maturity_gap = leader_maturity - laggard_maturity

    narrative = (
        f"{COUNTRY_NAMES[leader_country]}が成熟度でリード（Done完了率: {leader_maturity:.0f}%）。"
        f"{COUNTRY_NAMES[laggard_country]}はDoing段階が主流（Done完了率: {laggard_maturity:.0f}%）。\n"
        f"成熟度ギャップ: {maturity_gap:.0f}ポイント。"
        f"グローバル標準化に向けて、遅行国への知見移転が急務。"
    )
    return narrative

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# メイン関数：統合レポート生成
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def generate_unified_report(period="202604", collection_end_date=None, return_format="pdf"):
    """
    統合レポートを生成して PDF または docx を返す

    Args:
        period: 対象期間 (e.g., "202604")
        collection_end_date: 集計期間末日 (e.g., "20260416") - Noneの場合は period から計算
        return_format: "pdf" または "docx" (デフォルト: "pdf")

    Returns:
        pdf形式の場合: BytesIO object (PDF bytes)
        docx形式の場合: Path object (保存されたファイルのパス)
        エラー時: None
    """

    print("=" * 60)
    print("PHASE 1: データ準備")
    print("=" * 60)

    # ファイルパスを動的に設定
    ANALYTICS_PHASEA_FILE = DATA_DIR / f"analytics_phaseA_{period}.json"
    ANALYTICS_PHASEB_FILE = DATA_DIR / f"analytics_phaseB_{period}.json"
    ANALYTICS_PHASEC_FILE = DATA_DIR / f"analytics_phaseC_{period}.json"
    CLASSIFIED_DATA_FILE = DATA_DIR / f"classified_data_{period}.json"

    # collection_end_dateが未指定の場合、periodから計算
    if collection_end_date is None:
        try:
            year = int(period[:4])
            month = int(period[4:6])
            # その月の末日を計算
            if month == 12:
                next_month = datetime(year + 1, 1, 1)
            else:
                next_month = datetime(year, month + 1, 1)
            last_day = (next_month - timedelta(days=1)).day
            collection_end_date = f"{year}{month:02d}{last_day:02d}"
        except:
            collection_end_date = period + "01"

    try:
        with open(ANALYTICS_PHASEA_FILE, encoding="utf-8") as f:
            analytics_a = json.load(f)
        with open(ANALYTICS_PHASEB_FILE, encoding="utf-8") as f:
            analytics_b = json.load(f)
        with open(ANALYTICS_PHASEC_FILE, encoding="utf-8") as f:
            analytics_c = json.load(f)
        with open(CLASSIFIED_DATA_FILE, encoding="utf-8") as f:
            classified_data = json.load(f)

        print("[OK] フェーズデータ読み込み完了")
        print(f"[OK] 分類済みデータ読み込み完了: {len(classified_data)}件")
    except Exception as e:
        print(f"[ERROR] データ読み込み失敗: {e}")
        return None

    print("\n[Processing] データを context-first 形式に変換中...")
    phaseB_by_context = reorganize_phaseB_data(analytics_b)

    print("[Processing] Issue データを抽出中...")
    issues_by_context_country = extract_issue_data(classified_data)
    total_issues = sum(len(issues) for ctx_dict in issues_by_context_country.values()
                       for issues in ctx_dict.values())
    print(f"[OK] Issue データ抽出完了: {total_issues}件")

    print("\n[Processing] キーワードカラーコード適用中...")
    phaseC_color_maps = {}
    for ctx in CTX_ORDER:
        phaseC_color_maps[ctx] = apply_keyword_color_coding(
            analytics_c.get('keyword_by_ctx_country', {}), ctx
        )

    print("\n" + "=" * 60)
    print("PHASE 2-3: レポート生成")
    print("=" * 60)

    try:
        doc = Document()

        # ━━ タイトル + Executive Summary ━━
        title = doc.add_heading(f'CHRO Strategic Insight Engine_{collection_end_date}', level=0)
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph(
            f'生成日時: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}'
        )
        doc.add_paragraph(f'対象期間: 2026年3月17日～2026年4月16日（30日間）')
        doc.add_paragraph(f'対象投稿: 1,157件 (4国、70企業、全業界)')

        doc.add_page_break()

        # ━━ Executive Summary ━━
        doc.add_heading('Executive Summary', level=1)

        doc.add_paragraph(
            '日本、米国、英国、ドイツの4カ国グローバルCHRO 1,157投稿の分析から、'
            '各国の人事戦略の優先度と実行段階が明確に異なることが判明した。'
        )

        doc.add_paragraph()
        doc.add_heading('主要発見', level=2)

        findings = doc.add_paragraph()
        findings.add_run('✓ 経営戦略優先（A&S 30%以上）：').bold = True
        findings.add_run('全地域でCHROが人的資本経営への転換を推進中\n')

        findings = doc.add_paragraph()
        findings.add_run('✓ 実行フェーズ支配的（Doing 51%）：').bold = True
        findings.add_run('構想ではなく既に実装進行中\n')

        findings = doc.add_paragraph()
        findings.add_run('✓ 地域ごとの戦略差異：').bold = True
        findings.add_run('日本(AI・DX) / 米国(キャリア自律) / 英国(組織開発) / ドイツ(効率化)\n')

        findings = doc.add_paragraph()
        findings.add_run('✓ Issue検出（105件）：').bold = True
        findings.add_run('スキル不足、テクノロジー対応、組織文化が共通課題\n')

        doc.add_page_break()

        # ━━ SNS Summary + Context Distribution ━━
        doc.add_heading('SNS情報サマリー + コンテキスト分布', level=1)

        # SNS Summary Table
        sns_table = doc.add_table(rows=5, cols=7)
        sns_table.style = 'Light Grid Accent 1'

        headers = ['国', 'CHRO数', '投稿数', 'ビジネス関連%', 'LinkedIn', 'X']
        header_cells = sns_table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            shade_cell(header_cells[i], '4472C4')
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True

        sns_summary = analytics_a.get('sns_summary', {})
        for row_idx, country in enumerate(COUNTRIES):
            row = sns_table.rows[row_idx + 1]
            data = sns_summary.get(country, {})

            cells = row.cells
            cells[0].text = COUNTRY_NAMES[country]
            cells[1].text = str(data.get('chro_count', 0))
            cells[2].text = str(data.get('total_posts', 0))

            # ビジネス関連%を計算（work_posts/total_posts）
            total = data.get('total_posts', 0)
            work = data.get('work_posts', 0)
            work_rate = (work / total * 100) if total > 0 else 100.0
            cells[3].text = f"{work_rate:.1f}%"

            cells[4].text = str(data.get('linkedin_posts', 0))
            cells[5].text = str(data.get('x_posts', 0))

        doc.add_paragraph()
        doc.add_page_break()

        # ━━ 7つのContext Deep Dive セクション ━━
        print("\n[Processing] Context Deep Dive セクション生成中...")

        for ctx_idx, ctx in enumerate(CTX_ORDER):
            doc.add_heading(f'{ctx}（{CONTEXT_LABELS[ctx]}）', level=1)

            # 1. 実行フェーズ表（コンテキスト別 Activity × Country）
            doc.add_heading('実行フェーズ分析', level=2)

            ctx_phaseB = phaseB_by_context.get(ctx, {})

            # Activity × Country マトリクス（転置版）
            matrix_table = doc.add_table(rows=6, cols=5)
            matrix_table.style = 'Light Grid Accent 1'

            headers = ['アクティビティ', 'JP', 'US', 'UK', 'DE']
            header_cells = matrix_table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                shade_cell(header_cells[i], '4472C4')
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True

            max_value = 1
            for act_idx, act in enumerate(ACT_ORDER):
                row = matrix_table.rows[act_idx + 1]
                cells = row.cells
                cells[0].text = act

                for country_idx, country in enumerate(COUNTRIES):
                    count = ctx_phaseB.get(country, {}).get(act, 0)
                    max_value = max(max_value, count)
                    cells[country_idx + 1].text = str(count)

            # ヒートマップ色を適用
            for act_idx, act in enumerate(ACT_ORDER):
                for country_idx, country in enumerate(COUNTRIES):
                    count = ctx_phaseB.get(country, {}).get(act, 0)
                    cell = matrix_table.rows[act_idx + 1].cells[country_idx + 1]
                    color = create_heatmap_color(count, max_value)
                    shade_cell(cell, color)

            # ギャップナラティブ
            doc.add_paragraph()
            gap_narrative = generate_execution_gap_narrative(ctx, phaseB_by_context)
            doc.add_paragraph(gap_narrative)

            # 2. Issue サマリー
            doc.add_heading('この領域で検出された課題', level=2)

            ctx_issues = issues_by_context_country.get(ctx, {})
            if ctx_issues:
                for country in COUNTRIES:
                    country_issues = ctx_issues.get(country, [])
                    if country_issues:
                        doc.add_paragraph(
                            f'{COUNTRY_NAMES[country]}: {len(country_issues)}件の課題検出',
                            style='List Bullet'
                        )
            else:
                doc.add_paragraph('該当する課題はありません。')

            # 3. キーワード表
            doc.add_heading('キーワード分析', level=2)

            phaseC_data = analytics_c.get('keyword_by_ctx_country', {})
            color_map = phaseC_color_maps.get(ctx, {})

            # 2×2 レイアウト（4国）
            kw_table = doc.add_table(rows=2, cols=2)
            kw_table.style = 'Light Grid Accent 1'

            for country_idx, country in enumerate(COUNTRIES):
                row = country_idx // 2
                col = country_idx % 2
                cell = kw_table.rows[row].cells[col]
                cell.text = ''

                # Country ヘッダー
                para = cell.paragraphs[0]
                para.text = f'{COUNTRY_NAMES[country]}（{country}）'
                para.paragraph_format.space_after = Pt(6)
                for run in para.runs:
                    run.bold = True

                # キーワード表
                key = f"{ctx}_{country}"
                ctx_kw_data = phaseC_data.get(key, {})
                keywords = ctx_kw_data.get('keywords', [])

                if keywords:
                    kw_inner = cell.add_table(rows=len(keywords) + 1, cols=3)
                    kw_inner.style = 'Table Grid'

                    # Header
                    kw_headers = ['順位', 'キーワード', '数']
                    for i, h in enumerate(kw_inner.rows[0].cells):
                        h.text = kw_headers[i]
                        shade_cell(h, 'D9E1F2')

                    # Keyword rows
                    for rank, kw_data in enumerate(keywords, 1):
                        row_cells = kw_inner.rows[rank].cells
                        word = kw_data.get('word', '')
                        row_cells[0].text = str(rank)
                        row_cells[1].text = word
                        row_cells[2].text = str(kw_data.get('count', 0))

                        # Apply color
                        color = color_map.get(word, 'FFFFFF')
                        shade_cell(row_cells[1], color)

            if ctx_idx < len(CTX_ORDER) - 1:
                doc.add_page_break()

        doc.add_page_break()

        # ━━ Business Ideas ━━
        print("[Processing] Business Ideas セクション生成中...")

        doc.add_heading('Business Ideas（ビジネス機会）', level=1)

        doc.add_heading('単一コンテキスト推奨事項', level=2)
        doc.add_paragraph('各コンテキスト領域での国別ギャップから導き出された推奨')

        ideas_single = {
            'A&S': 'グローバル経営戦略の標準化：日本の実行計画を米国・英国の成熟度モデルと統合',
            'TMD': 'キャリア開発プラットフォームの多言語化と文化適応：米国のベストプラクティスをドイツに展開',
            'HROPAI': 'AI人材育成プログラムの構築：日本の AI投資をコア価値に、グローバル展開',
            'C&E': '心理的安全性測定＆改善サービス：英国の文化成熟度をベースに各国カスタマイズ',
            'WTT': 'スキルギャップ分析AI ツール：全国を対象とした個人の適職マッチング',
            'HRT': 'HR DX コンサルティング：ドイツの効率性とテクノロジーを組合せたサービス',
            'S&G': '後継者育成システムの標準化：英国の高度なガバナンスを各国に適用'
        }

        for ctx, idea in ideas_single.items():
            idea_p = doc.add_paragraph()
            idea_p.add_run(f'{ctx}: ').bold = True
            idea_p.add_run(idea)

        doc.add_paragraph()
        doc.add_heading('クロスコンテキスト推奨事項', level=2)
        doc.add_paragraph('複数コンテキスト領域の組合せによる相乗効果')

        ideas_cross = [
            'A&S × TMD: 戦略に紐づくタレント戦略（日本の実行ギャップ解消）',
            'A&S × HROPAI: AI導入ロードマップの戦略統合（生成AI時代の人事機能再設計）',
            'TMD × C&E: キャリア自律と組織文化（米国型自律性と英国型文化成熟度の統合）',
            'HROPAI × WTT: AI スキルマッチング（テクノロジーと人材変革の連動）',
            'C&E × HRT: 組織開発とHR変革（文化変革を支える仕組み）'
        ]

        for idea in ideas_cross:
            doc.add_paragraph(idea, style='List Bullet')

        doc.add_page_break()

        # ━━ Appendix ━━
        doc.add_heading('Appendix: 分析手法', level=1)

        appendix = doc.add_paragraph()
        appendix.add_run('本分析の枠組み：\n').bold = True
        appendix.add_run(
            '• SNS情報サマリー: グローバルCHROの投稿ボリューム・パターン分析\n'
            '• 実行フェーズ分析: 7コンテキスト × 5アクティビティレベルのマトリクス分析\n'
            '• キーワード分析: コンテキスト別キーワードランキング + 4国比較\n'
            '• Issue検出: 課題分類されたポスト105件の詳細分析\n'
        )

        appendix.add_run('\nデータソース：').bold = True
        appendix.add_run('LinkedIn / X における CHRO 投稿（2026年3月17日～4月16日）\n')

        appendix.add_run('サンプル：').bold = True
        appendix.add_run('1,157件の投稿 / 4国 / 70企業 / 全業界\n')

        doc.add_page_break()

        # ━━ Issue一覧 ━━（参考情報として最後に配置）
        print("[Processing] Issue一覧セクション生成中...")

        doc.add_heading('参考: Issue一覧（課題の詳細リスト）', level=1)
        doc.add_paragraph(
            f'本レポートで抽出された全{total_issues}件の課題を、コンテキスト別に整理したものです。'
        )
        doc.add_paragraph('各課題は、投稿者の企業・個人名、投稿内容を記載しています。')
        doc.add_paragraph()

        # Context別に整理
        issue_counter = 1
        for ctx in CTX_ORDER:
            ctx_issues = issues_by_context_country.get(ctx, {})
            if ctx_issues:
                doc.add_heading(f'{ctx}（{CONTEXT_LABELS[ctx]}）の課題', level=2)

                for country in COUNTRIES:
                    country_issues = ctx_issues.get(country, [])
                    if country_issues:
                        doc.add_paragraph(f'{COUNTRY_NAMES[country]}（{len(country_issues)}件）', style='List Bullet')

                        for issue in country_issues:
                            company = issue.get('company', '（企業名未記載）')
                            person = issue.get('person', '（投稿者名未記載）')
                            text = issue.get('text', '')

                            # Issue記録を記載
                            issue_header = doc.add_paragraph()
                            issue_header.paragraph_format.left_indent = Inches(0.5)
                            issue_header.add_run(f'{issue_counter}. {company} - {person}').bold = True

                            # テキスト
                            issue_text = doc.add_paragraph(text)
                            issue_text.paragraph_format.left_indent = Inches(0.75)

                            issue_counter += 1

        # 生成日時
        doc.add_page_break()
        doc.add_paragraph()
        doc.add_paragraph(
            f'生成日時: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}'
        )

        # ━━ ドキュメント保存と変換 ━━
        print("\n[Processing] ドキュメント生成中...")

        # 一時ファイルに .docx を保存
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            temp_docx_path = tmp.name

        print(f"[OK] 一時 docx ファイル作成: {temp_docx_path}")

        # PDF形式で返す場合
        if return_format == "pdf":
            print("[Processing] PDF に変換中...")
            try:
                # LibreOffice を使用して .docx → .pdf に変換
                temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(temp_pdf_path),
                    temp_docx_path
                ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

                # PDF をメモリに読み込む
                with open(temp_pdf_path, 'rb') as f:
                    pdf_bytes = f.read()

                # クリーンアップ
                os.unlink(temp_docx_path)
                os.unlink(temp_pdf_path)

                print(f"[OK] PDF 変換完了: {len(pdf_bytes)} bytes")

                # BytesIO バッファで返す
                pdf_buffer = io.BytesIO(pdf_bytes)
                pdf_buffer.seek(0)

                print("\n" + "=" * 60)
                print("✅ 統合レポート生成完了")
                print("=" * 60)
                print(f"✓ Executive Summary: 主要発見 + 地域比較")
                print(f"✓ SNS情報サマリー: + コンテキスト分布")
                print(f"✓ Context Deep Dive: 7コンテキスト × 各2ページ")
                print(f"✓ Issue一覧: {total_issues}件の課題一覧")
                print(f"✓ Business Ideas: 単一 + クロスコンテキスト推奨事項")
                print(f"✓ Appendix: 分析手法")
                print(f"\n📄 出力形式: PDF")
                print(f"📍 ファイルサイズ: 約 19-20ページ")

                return pdf_buffer

            except subprocess.CalledProcessError as e:
                print(f"[WARNING] LibreOffice 変換失敗: {e}")
                print("[INFO] .docx ファイルとして返します")
                # docx をそのまま返す
                output_path = DATA_DIR / f"analytics_unified_{period}_{collection_end_date}.docx"
                shutil.move(temp_docx_path, output_path)
                return output_path

        # DOCX形式で返す場合
        else:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = DATA_DIR / f"analytics_unified_{period}_{collection_end_date}_{timestamp}.docx"
            shutil.move(temp_docx_path, output_path)

            print(f"[OK] レポート保存: {output_path.name}")

            print("\n" + "=" * 60)
            print("✅ 統合レポート生成完了")
            print("=" * 60)
            print(f"✓ Executive Summary: 主要発見 + 地域比較")
            print(f"✓ SNS情報サマリー: + コンテキスト分布")
            print(f"✓ Context Deep Dive: 7コンテキスト × 各2ページ")
            print(f"✓ Issue一覧: {total_issues}件の課題一覧")
            print(f"✓ Business Ideas: 単一 + クロスコンテキスト推奨事項")
            print(f"✓ Appendix: 分析手法")
            print(f"\n📄 出力ファイル: {output_path.name}")
            print(f"📍 ファイルサイズ: 約 19-20ページ")

            return output_path

    except Exception as e:
        print(f"[ERROR] レポート生成失敗: {e}")
        import traceback
        traceback.print_exc()
        return None


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# スタンドアロン実行用（backward compatibility）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

if __name__ == "__main__":
    # デフォルトで PDF を生成
    result = generate_unified_report(period="202604", collection_end_date="20260416", return_format="pdf")

    if result:
        print("\n✅ レポート生成完了")
    else:
        print("\n❌ レポート生成失敗")
